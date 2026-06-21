"""DOCX file parser."""
import logging
from .base import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> dict:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        text = "\n".join(text_parts)

        # Extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)

        return {"text": text, "tables": tables, "language": self._detect_language(text)}

    def _detect_language(self, text: str) -> str:
        import re
        ru_chars = len(re.findall(r'[а-яА-ЯёЁ]', text))
        en_chars = len(re.findall(r'[a-zA-Z]', text))
        return "ru" if ru_chars > en_chars else "en"
